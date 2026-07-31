from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "scripts"))

from lib.docx_extract import extract_docx_text
from lib.latex_render import render_latex
from lib.resume_data import (
    load_schema,
    load_yaml,
    schema_errors,
    semantic_errors,
)

EXAMPLE = REPO_ROOT / "examples/fictional-resume.yaml"
SCHEMA = REPO_ROOT / "data/schema/resume.schema.json"


def test_fictional_example_is_valid() -> None:
    data = load_yaml(EXAMPLE)
    assert schema_errors(data, load_schema(SCHEMA)) == []
    assert semantic_errors(data) == []


def test_unknown_metric_reference_is_rejected() -> None:
    data = copy.deepcopy(load_yaml(EXAMPLE))
    statement = data["sections"]["experience"][0]["highlights"][0]
    statement["metric_refs"].append("missing_metric")
    errors = semantic_errors(data)
    assert any("unknown metrics missing_metric" in error for error in errors)


def test_render_ready_rejects_unconfirmed_claims() -> None:
    data = copy.deepcopy(load_yaml(EXAMPLE))
    data["sections"]["projects"][0]["verification"]["status"] = "needs_confirmation"
    errors = semantic_errors(data)
    assert any("render-ready" in error for error in errors)


def test_latex_rendering_escapes_special_characters() -> None:
    data = copy.deepcopy(load_yaml(EXAMPLE))
    data["basics"]["name"]["en"] = "Avery & Lin"
    data["sections"]["experience"][0]["organization"]["en"] = "R&D_Studio"
    rendered = render_latex(data, "en", REPO_ROOT / "template")
    assert r"Avery \& Lin" in rendered
    assert r"R\&D\_Studio" in rendered
    assert r"\documentclass[english]{cv}" in rendered


def test_optional_locations_and_project_dates_can_be_omitted() -> None:
    data = copy.deepcopy(load_yaml(EXAMPLE))
    data["basics"].pop("location")
    for section in ("education", "experience", "activities"):
        for item in data["sections"][section]:
            item.pop("location", None)
    for project in data["sections"]["projects"]:
        project.pop("date", None)

    assert schema_errors(data, load_schema(SCHEMA)) == []
    assert semantic_errors(data) == []
    for locale in ("en", "zh"):
        rendered = render_latex(data, locale, REPO_ROOT / "template")
        assert r"\begin{document}" in rendered


def test_docx_reference_extraction_includes_layout_regions() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Candidate Header"
    document.add_paragraph("Education")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Institution"
    table.cell(0, 1).text = "2024"
    document.sections[0].footer.paragraphs[0].text = "Private Footer"

    extracted = extract_docx_text(document)
    assert "[header]" in extracted
    assert "Candidate Header" in extracted
    assert "[body]" in extracted
    assert "Institution | 2024" in extracted
    assert "[footer]" in extracted
